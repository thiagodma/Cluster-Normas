#Importando os pacotes que seraoo utilizados
from stop_words import get_stop_words
from docx import Document
import os, os.path, glob, re, unicodedata, time, nltk
from sklearn.decomposition import TruncatedSVD
from wordcloud import WordCloud
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.cm as cm
import pdb
#nltk.download('rslp')
#nltk.download('stopwords')


class ClusterNormas:
    """Essa classe contem todos os metodos e atributos necessarios para fazer a
    clusterização das normas da ANVISA utilizando-se de uma tecnica classica."""

    def __init__(self):
        '''Construtor da classe ClusterNormas'''
        self.tipos_norma = {'PRT':'Portaria', 'RDC':'RDC', 'RES':'RE', 'RE':'RE',
                       'IN':'IN', 'INC':'INC', 'PRTC':'PRTC'}
        self.resolucoes = []
        self.nome_arquivos= []
        self.macrotemas = pd.read_csv('Macrotemas_python.csv',sep='|').rename(columns={'Macrotema atual':'Macrotema_atual'}).Macrotema_atual.unique()
        self.macrotema_por_norma = []
        self.tabela_macrotemas = pd.read_csv('Macrotemas_python.csv',sep='|').rename(columns={'Macrotema atual':'Macrotema_atual'})
        self.df_resolucoes_macrotemas = []


    def trata_textos(self, texto):

        texto = re.sub(r'\xa0',' ',texto)

        #encontra a seção de artigos
        artigos = re.findall(r'\n *Art. *\d',texto)

        if len(artigos) >=2:
            #Pega apenas o texto entre o primeiro artigo e o último
            regex = r'('+artigos[0]+')(.*)('+artigos[-1]+')'
            regex = re.sub(r'\n',r'\\n',regex)
            m = re.search(regex, texto, re.DOTALL)
            texto_so_artigos = m.group(2)
        else:
            return 'norma fora de padrão'

        texto_limpo = 'Art 1'+texto_so_artigos

        return texto_limpo

    def escreve_artigos_em_txt(self):

        macrotemas_aux = [macrotema.lower() for macrotema in self.macrotemas]
        self.macrotemas = list(dict.fromkeys(macrotemas_aux))

        macrotema_por_norma_aux = [macrotema.lower() for macrotema in self.macrotema_por_norma]
        self.macrotema_por_norma = macrotema_por_norma_aux

        if not os.path.isdir('Data'): os.mkdir('Data')

        os.chdir('Data')

        for macrotema in self.macrotemas:
            if not os.path.isdir(macrotema): os.mkdir(macrotema)
            os.chdir(macrotema)
            for i in range(len(self.resolucoes)):
                if self.macrotema_por_norma[i] == macrotema:
                    fo = open(self.nome_arquivos[i][:-5]  + '.txt', 'w+')
                    fo.writelines(self.resolucoes[i])
                    fo.close()
            os.chdir('..')
        os.chdir('..')


    def normaliza_nome_arquivo(self, nome_arq):
        '''Esta funcao serve para padronizar os nomes dos arquivos'''

        n_tratado = nome_arq
        # Procura o ultimo digito do nome do arquivo
        acha_digitos = re.match('.+([0-9])[^0-9]*$', nome_arq)
        # Posicao do ultimo digito encontrada
        pos_ultimo_digito = acha_digitos.start(1)
        # Corta o resto alem do ultimo di­gito
        nome_arq = nome_arq[:pos_ultimo_digito + 1]
        # Separa no split por underline
        nome_arq = nome_arq.split('_')
        # Parada de emergencia para nome de arquivos fora do padrao
        if len(nome_arq) < 3:
            raise ValueError('Nome de arquivo fora de padrao:', n_tratado)
        # Monta nome_arq em norma_citadora no padrao desejado tipo_norma numero_norma/AAAA
        # JÃ¡ Ã© aplicado tambem o dicionario tipos_norma ao nome_arq[0], normalizando
        # assim o tipo da norma
        numeracao = nome_arq[-2].strip()
        # Adicionando o '0' Ã s normas que sao numeral unico, como IN 1/2014
        if len(numeracao) == 1: numeracao = '0' + numeracao
        norma_citadora = self.tipos_norma[nome_arq[0].strip()] + ' ' + numeracao + '_' + nome_arq[-1]
        # Retorna nome_arq normalizado
        return norma_citadora

    def identifica_macrotema(self,nome_arquivo):
        '''
        Identifica a qual macrotema uma norma pertence. Faz isso consultando um csv.
        '''

        #carrega o dataframe que contém a informação do macrotema de cada norma
        df = self.tabela_macrotemas

        #Modifica o nome do arquivo para ficar no mesmo padrão do csv
        nome_arquivo = re.sub('_','/',nome_arquivo[:-5])

        #identifica o indice da norma 'nome_arquivo' no dataframe
        idx = df[df['Citadora'] == nome_arquivo].index.tolist()

        #Se achou a norma na tabela
        if len(idx) != 0:
            macrotema = df['Macrotema_atual'].iloc[idx[0]]
        else:
            macrotema = None
            #import pdb; pdb.set_trace()

        return macrotema


    def importa_normas(self):
        '''
        Importa arquivos. Esse código deve funcionar em qualquer computador em que a pasta 'Arquivos_DOCX_atual_31_julho_2019'
        esteja no diretório de trabalho
        '''

        path = 'Arquivos_DOCX_atual_31_julho_2019'
        os.chdir(path)

        # Lista de pastas no diretorio atual
        pastas = [name for name in os.listdir('.')]

        arquivos = []

        for i in range(0,len(pastas)):
            # Faz o diretorio mudar para a pasta de cada ano
            os.chdir(pastas[i])
            # Pega todos os arquivos da pasta
            arquivos = glob.glob('*.docx')
            print('\nA pasta ' + str(i) + ' tem ' + str(len(arquivos)) + ' arquivos')
            t = time.time()
            for w in range(0, len(arquivos)):
                if int(arquivos[w][-9:-5]) >= 1999: #pega apenas as normas de 1999 pra frente
                    doc = Document(arquivos[w])
                    doc.paragraphs
                    texto = [parag.text for parag in doc.paragraphs]
                    texto = '\n'.join(texto)
                    # Substituicoes para desonerar o vetor Start, legado
                    self.resolucoes.append(self.trata_textos(texto))
                    self.nome_arquivos.append(arquivos[w])
                    macrotema = self.identifica_macrotema(arquivos[w])
                    if macrotema != None:
                        self.macrotema_por_norma.append(macrotema)

            elapsed = time.time() - t
            print('Tempo para importar a pasta ' + str(i) + ': ' + str(elapsed) + '\n')
            os.chdir('..')
        #volta para o diretorio inicial
        os.chdir('..')


    def analisa_clusters(self, base_tfidf, id_clusters, macrotema):
        '''
        Tenta visualizar as cluster definidas. Além disso retorna um dataframe
        que contem a informacao do numero de normas por cluster
        '''

        clusters = np.unique(id_clusters)

        #inicializa o output da funcao
        n_normas = np.zeros(len(clusters)) #numero de normas pertencentes a uma cluster

        for cluster in clusters:
            idxs = np.where(id_clusters == cluster) #a primeira cluster não é a 0 e sim a 1
            n_normas[cluster-1] = len(idxs[0])

        macrotema_list = [macrotema]*len(clusters)
        macrotema_cluster_nnormas = pd.DataFrame(list(zip(macrotema_list,clusters,n_normas)),columns=['macrotema','cluster_id','n_normas'])

        return macrotema_cluster_nnormas


    def generate_wordcloud(self, cluster):
        '''
        Gera uma nuvem de palavras de uma cluster 'cluster'.
        '''

        df = pd.read_csv('cluster_normas_cosseno.csv',sep='|')
        a = df[df['cluster_id'] == cluster]

        L=[]
        for i in range(a.shape[0]):
            L.append(self.resolucoes_tratadas[a.iloc[i,2]])

        text = '\n'.join(L)


        wordcloud = WordCloud(stopwords=self.stop_words.split()+['laboratorio','laboratorios','rdc']).generate(text)
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        plt.show()
